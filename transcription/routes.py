import os
import time
import uuid
import threading
import jwt
from flask import Blueprint, request, Response, jsonify, render_template, send_from_directory, session
from config import Config
from transcription.service import run_transcription, processing_jobs, allowed_file
from deep_translator import GoogleTranslator
from docx import Document
from io import BytesIO
from auth.utils import token_required
from werkzeug.utils import secure_filename


# ============================================================
# BLUEPRINT (MUST BE FIRST)
# ============================================================
transcription_bp = Blueprint("transcription", __name__)


# ============================================================
# HOME PAGE
# ============================================================
@transcription_bp.route('/')
def index():
    files = []
    upload_path = Config.UPLOAD_FOLDER

    if os.path.exists(upload_path):
        for f in os.listdir(upload_path):
            if f.lower().endswith(('mp3', 'wav', 'm4a', 'flac', 'aac', 'ogg')):
                full_path = os.path.join(upload_path, f)
                files.append({
                    'name': f,
                    'time': time.strftime(
                        "%Y-%m-%d %H:%M:%S",
                        time.localtime(os.path.getctime(full_path))
                    )
                })

    files.sort(key=lambda x: x['time'], reverse=True)

    return render_template('index.html', files=files)


# ============================================================
# SERVE AUDIO
# ============================================================
@transcription_bp.route('/audio/<filename>')
def serve_audio(filename):
    return send_from_directory(Config.UPLOAD_FOLDER, filename)


# ============================================================
# UPLOAD
# ============================================================
@transcription_bp.route('/upload', methods=['POST'])
def upload_file():

    # Initialize trial counter
    auth_header = request.headers.get("Authorization")

    is_logged_in = False

    if auth_header:
        try:
            token = auth_header.split()[1]
            jwt.decode(token, Config.SECRET_KEY, algorithms=[Config.ALGORITHM])
            is_logged_in = True
        except:
            pass

    if not is_logged_in:
        if "trial_count" not in session:
            session["trial_count"] = 0

        if session["trial_count"] >= 2:
            return jsonify({
                "trial_ended": True,
                "message": "Free Trial Ended. Please login to continue."
            }), 403


    if 'audio' not in request.files:
        return jsonify({"error": "No file"}), 400

    file = request.files['audio']

    if file and allowed_file(file.filename):

        original_filename = secure_filename(file.filename)

        name_part, ext_part = os.path.splitext(original_filename)

        unique_suffix = f"{int(time.time())}_{uuid.uuid4().hex[:6]}"

        filename = f"{name_part}_{unique_suffix}{ext_part}"

        file_path = os.path.join(Config.UPLOAD_FOLDER, filename)
        file.save(file_path)

        # Increase trial count ONLY ONCE
        if "user_id" not in session:
            session["trial_count"] += 1

        threading.Thread(
            target=run_transcription,
            args=(file_path, filename)
        ).start()

        return jsonify({"message": "Upload successful", "filename": filename})

    return jsonify({"error": "Invalid file"}), 400

# ============================================================
# CHECK STATUS (Polling)
# ============================================================
@transcription_bp.route('/check_status/<filename>')
def check_status(filename):
    return jsonify(
        processing_jobs.get(
            filename,
            {"status": "processing", "progress": 5}
        )
    )


# ============================================================
# DOWNLOAD (TXT / PDF)
# ============================================================
@transcription_bp.route('/download/<filename>')
@token_required
def download(filename):

    file_type = request.args.get('type', 'txt')
    target_lang = request.args.get('lang', 'en')

    transcript_path = os.path.join(Config.UPLOAD_FOLDER, filename + ".txt")
    summary_path = os.path.join(Config.UPLOAD_FOLDER, filename + "_summary.txt")

    if not os.path.exists(transcript_path):
        return "File not found", 404

    with open(transcript_path, "r", encoding="utf-8") as f:
        transcript = f.read()

    summary = ""
    if os.path.exists(summary_path):
        with open(summary_path, "r", encoding="utf-8") as f:
            summary = f.read()

    # Translate if needed
    if target_lang != "en":
        translator = GoogleTranslator(source="auto", target=target_lang)

        transcript_chunks = [
            transcript[i:i+4500]
            for i in range(0, len(transcript), 4500)
        ]

        transcript = "".join(
            translator.translate(chunk)
            for chunk in transcript_chunks if chunk.strip()
        )

        if summary:
            summary = translator.translate(summary)

    # TXT
    if file_type == "txt":
        content = f"SUMMARY:\n{summary}\n\nTRANSCRIPT:\n{transcript}"

        return Response(
            content.encode("utf-8"),
            mimetype="text/plain",
            headers={
                "Content-Disposition": f"attachment;filename={filename}_{target_lang}.txt"
            }
        )

    # DOCX
    if file_type == "docx":
        document = Document()
        document.add_heading("TranscribeFlow Report", level=1)

        document.add_heading("Summary", level=2)
        document.add_paragraph(summary if summary else "No summary available.")

        document.add_heading("Transcript", level=2)
        document.add_paragraph(transcript)

        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)

        return Response(
            file_stream.getvalue(),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment;filename={filename}_{target_lang}.docx"
            }
        )
    return "Invalid file type", 400

# ============================================================
# DELETE SINGLE FILE
# ============================================================
@transcription_bp.route('/delete/<filename>', methods=['POST'])
@token_required
def delete_file(filename):
    try:
        file_path = os.path.join(Config.UPLOAD_FOLDER, filename)

        if os.path.exists(file_path):
            os.remove(file_path)

        for ext in [".txt", "_summary.txt"]:
            extra_file = os.path.join(Config.UPLOAD_FOLDER, filename + ext)
            if os.path.exists(extra_file):
                os.remove(extra_file)

        processing_jobs.pop(filename, None)

        return jsonify({"success": True})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============================================================
# CLEAR ALL (AJAX)
# ============================================================
@transcription_bp.route('/clear_all', methods=['POST'])
@token_required
def clear_all():
    try:
        for f in os.listdir(Config.UPLOAD_FOLDER):
            file_path = os.path.join(Config.UPLOAD_FOLDER, f)
            if os.path.isfile(file_path):
                os.remove(file_path)

        processing_jobs.clear()

        return jsonify({"success": True})

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ============================================================
# TRANSLATE ON FLY
# ============================================================
@transcription_bp.route('/translate_on_fly', methods=['POST'])
def translate_on_fly():
    data = request.get_json()

    transcript = data.get('transcript', '')
    summary = data.get('summary', '')
    target_lang = data.get('target', 'en')

    try:
        translator = GoogleTranslator(source='auto', target=target_lang)

        translated_text = ""
        for chunk in [
            transcript[i:i+4500]
            for i in range(0, len(transcript), 4500)
        ]:
            translated_text += translator.translate(chunk)

        translated_summary = translator.translate(summary) if summary else ""

        return jsonify({
            'success': True,
            'translated_text': translated_text,
            'translated_summary': translated_summary
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500
